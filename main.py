from openpyxl import load_workbook
from docx import Document

# Load the workbook
excel_path = './working_files/budget.xlsx'
workbook = load_workbook(filename=excel_path)
sheet = workbook.active

# Create a new Word document# Load the workbook
excel_path = './working_files/budget.xlsx'
workbook = load_workbook(filename=excel_path, data_only=True)  # Add data_only=True
sheet = workbook.active
doc = Document()

# Iterate through each row in the Excel file
for row in sheet.iter_rows(min_row=2, values_only=True):
    project_title, project_description, business_driver, business_value, business_risk, \
    budget_expense, internal_hours, external_hours, solutions_involvement, \
    solutions_hours, pmo_involvement, pmo_hours, total_expected_hours = row

    # Add data to the Word document while making the title bold
    p= doc.add_paragraph()
    p.add_run('Project Title').bold = True
    p.add_run(f': {project_title}')
    p = doc.add_paragraph()
    p.add_run('Project Description').bold = True
    p.add_run(f': {project_description}')
    p = doc.add_paragraph()
    p.add_run('Business Driver').bold = True
    p.add_run(f': {business_driver}')
    p = doc.add_paragraph()
    p.add_run('Business Value').bold = True
    p.add_run(f': {business_value}')
    p = doc.add_paragraph()
    p.add_run('Business Risk').bold = True
    p.add_run(f': {business_risk}')
    p = doc.add_paragraph()
    p.add_run('Budget Expense').bold = True
    p.add_run(f': {budget_expense}')
    p = doc.add_paragraph()
    p.add_run('Internal Hours').bold = True
    p.add_run(f': {internal_hours}')
    p = doc.add_paragraph()
    p.add_run('External Hours').bold = True
    p.add_run(f': {external_hours}')
    p = doc.add_paragraph()
    p.add_run('Solutions Involvement').bold = True
    p.add_run(f': {solutions_involvement}')
    p = doc.add_paragraph()
    p.add_run('Solutions Hours').bold = True
    p.add_run(f': {solutions_hours}')
    p = doc.add_paragraph()
    p.add_run('PMO Involvement').bold = True
    p.add_run(f': {pmo_involvement}')
    p = doc.add_paragraph()
    p.add_run('PMO Hours').bold = True
    p.add_run(f': {pmo_hours}')
    p = doc.add_paragraph()
    p.add_run('Total Expected Hours').bold = True
    p.add_run(f': {total_expected_hours}')

    # Add a page break after each project
    doc.add_page_break()

# Save the Word document
word_document_path = './working_files/budget.docx'
doc.save(word_document_path)